from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path("output/doc")
OUTPUT_PATH = OUTPUT_DIR / "paper_plane_prd_formal.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10.5)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(33, 37, 41)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11.5)
    run.bold = True


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = p.add_run(bold_prefix)
        prefix_run.bold = True
        prefix_run.font.name = "Microsoft YaHei"
        prefix_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        prefix_run.font.size = Pt(10.5)
        rest_run = p.add_run(text[len(bold_prefix):])
        rest_run.font.name = "Microsoft YaHei"
        rest_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        rest_run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "项目", bold=True)
    set_cell_text(hdr[1], "内容", bold=True)
    shade_cell(hdr[0], "D9EAF7")
    shade_cell(hdr[1], "D9EAF7")
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(11)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        shade_cell(header_cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    doc.add_paragraph()
    return table


def build_doc():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("《纸飞机冲榜》正式PRD文档")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(26, 86, 132)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    subrun = subtitle.add_run("网页端轻休闲物理投掷竞技游戏产品需求文档")
    subrun.font.name = "Microsoft YaHei"
    subrun._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subrun.font.size = Pt(12)

    add_two_col_table(
        doc,
        [
            ("项目名称", "纸飞机冲榜"),
            ("文档类型", "正式PRD"),
            ("适用平台", "网页端"),
            ("版本号", "V1.0"),
            ("文档日期", datetime.now().strftime("%Y-%m-%d")),
            ("文档目的", "用于产品、设计、前端、后端与测试团队对齐首版目标、玩法与交付范围"),
        ],
    )

    add_heading(doc, "1. 产品概述", level=1)
    add_paragraph(doc, "《纸飞机冲榜》是一款网页端轻休闲 + 物理投掷 + 成绩竞争游戏。玩家通过选择纸张、调整参数化折纸结构，并完成一次投掷，让纸飞机飞出尽可能远的距离，以刷新个人纪录并冲击排行榜。")
    add_paragraph(doc, "产品核心一句话：玩家在网页上折出一架纸飞机，通过调整结构和投掷力度，让它飞出尽可能远的距离，并冲击排行榜。", bold_prefix="产品核心一句话：")

    add_heading(doc, "2. 立项背景与产品目标", level=1)
    add_paragraph(doc, "本项目面向网页小游戏场景，强调短时可玩、规则易懂、反馈直接与可重复挑战。纸飞机题材天然具备低门槛认知和物理想象空间，适合作为轻竞技玩法的包装载体。")
    add_bullets(
        doc,
        [
            "目标一：让玩家在10秒内理解玩法并完成首次构筑。",
            "目标二：让玩家在前3至5局中明确感知不同参数组合会带来不同飞行结果。",
            "目标三：通过个人最佳与总榜竞争提升重复游玩意愿。",
            "目标四：控制首版实现复杂度，优先验证玩法手感、可解释性与排行榜驱动力。",
        ],
    )

    add_heading(doc, "3. 用户定位与使用场景", level=1)
    add_table(
        doc,
        ["维度", "内容"],
        [
            ("核心用户", "喜欢轻量网页小游戏、愿意反复刷新成绩的泛休闲玩家"),
            ("次级用户", "喜欢研究最优解、比较配置效果的数值型玩家"),
            ("典型场景", "碎片时间、社交分享、办公室或课堂间隙短时游玩"),
            ("单局时长", "20至40秒"),
            ("核心动机", "冲更远距离、找到更优折法、进入排行榜"),
        ],
    )

    add_heading(doc, "4. 产品定位与卖点", level=1)
    add_bullets(
        doc,
        [
            "低门槛：玩家无需理解真实空气动力学，也能快速上手。",
            "高反馈：每次调参、投掷、飞行都能直观看到结果变化。",
            "强优化感：失败后玩家能归因到纸张、参数或投掷动作，而不是纯随机。",
            "轻竞技：排行榜使玩法具备长期挑战价值和传播价值。",
        ],
    )

    add_heading(doc, "5. 核心体验与设计原则", level=1)
    add_numbered(
        doc,
        [
            "可理解：玩家在第一次游玩时就能建立基础因果关系，如头更尖更快、翼更宽更能飘。",
            "可调优：每次失败都能促使玩家尝试新的纸张、参数或投掷角度。",
            "可对比：结算阶段必须清晰呈现本次表现与历史最佳、榜单门槛之间的差距。",
            "可复玩：单局时长短、重开成本低、优化空间持续存在。",
        ],
    )

    add_heading(doc, "6. 核心循环", level=1)
    add_heading(doc, "6.1 单局循环", level=2)
    add_paragraph(doc, "进入游戏 -> 选择纸张与调整参数 -> 完成折纸动画 -> 投掷起飞 -> 自动飞行与轻微扰动 -> 结算距离 -> 判断是否上榜 -> 再来一局。")
    add_heading(doc, "6.2 长线循环", level=2)
    add_paragraph(doc, "玩家通过反复尝试，不断优化折法与投掷方式，逐步提升个人最远纪录，并以进入更高榜单排名为中长期目标。")

    add_heading(doc, "7. 玩法详细设计", level=1)
    add_heading(doc, "7.1 参数化折纸系统", level=2)
    add_paragraph(doc, "本作不采用真实折纸模拟，而是使用半拟真折纸演出包装底层数值调整。玩家在视觉上完成对折、折机头、折机翼、调尾翼四个步骤，实际是在配置四个关键飞行参数。")
    add_table(
        doc,
        ["参数", "范围", "玩家认知", "正向收益", "风险与代价"],
        [
            ("机头尖锐度", "0至100", "头更尖，飞得更快", "初速度提升，阻力下降", "稳定性下降，俯冲倾向增强"),
            ("机翼宽度", "0至100", "翼更宽，飞得更久", "升力提升，滑翔时间增长", "初速度下降，更容易受扰动"),
            ("左右对称度", "基础值+微调", "更对称，更稳定", "偏航减少，成绩波动变小", "上限收益不一定最高"),
            ("尾翼上翘角度", "0至100", "尾翼决定抬头还是下坠", "更容易抬头滑翔", "过高可能失速，过低会快速下坠"),
        ],
    )
    add_paragraph(doc, "交互要求：每完成一步折纸，右侧实时刷新性能预估条，包括速度、升力、稳定性和风险提示。参数反馈必须是即时的，避免玩家在结算前无法理解自己做了什么。")

    add_heading(doc, "7.2 纸张系统", level=2)
    add_paragraph(doc, "纸张系统承担机体底盘差异，是首版实现流派多样性的核心手段之一。纸张不是纯皮肤，而是会改变飞机的基础物性。")
    add_table(
        doc,
        ["纸张类型", "定位", "数值特征", "推荐玩法"],
        [
            ("普通打印纸", "均衡型", "重量中、稳定中、升力中", "新手首选，适合稳定发挥"),
            ("轻薄笔记纸", "滑翔型", "重量低、升力高、抗扰动低", "适合冲击高上限远距离"),
            ("硬质卡纸", "冲刺型", "重量高、速度高、稳定高、滑翔弱", "适合高速开局和保守发挥"),
        ],
    )

    add_heading(doc, "7.3 投掷系统", level=2)
    add_paragraph(doc, "投掷是本作的第二层技术点，用于放大操作差异。折纸决定纸飞机素质，投掷决定本局发挥。")
    add_table(
        doc,
        ["变量", "输入方式", "影响", "设计目标"],
        [
            ("力度", "按住蓄力", "影响初速度和失控风险", "避免无脑拉满，强调时机判断"),
            ("角度", "向上拖动", "影响起飞高度和中段滑翔空间", "让玩家掌握适宜角度区间"),
            ("稳定释放", "松手时机", "影响起飞抖动与稳定加成", "形成可练习的操作上限"),
        ],
    )
    add_paragraph(doc, "建议设计一个最佳释放区。当玩家在最佳区间松手时，获得额外稳定性加成；过早或过晚则导致起飞抖动加剧，提升成绩差异。")

    add_heading(doc, "7.4 飞行系统", level=2)
    add_paragraph(doc, "飞行阶段以自动演出为主，不设置中途操控。玩家关注的是飞行轨迹、距离增长与最终落点，而不是持续操作。")
    add_bullets(
        doc,
        [
            "起飞段：由初速度、投掷角度和机头形状决定冲刺表现。",
            "滑翔段：由升力、尾翼角度和纸张重量决定滞空时长。",
            "下坠段：由重量、阻力和稳定性决定落地前的距离收益。",
            "轻微扰动：每局存在少量随机空气扰动，稳定性越低，受影响越明显。",
        ],
    )

    add_heading(doc, "8. 游戏规则与结算", level=1)
    add_heading(doc, "8.1 结算指标", level=2)
    add_bullets(
        doc,
        [
            "本次飞行距离",
            "是否刷新个人最佳",
            "是否进入排行榜",
            "本局飞机配置摘要",
            "失败原因提示，例如俯冲过早、失速明显、偏航过大",
        ],
    )
    add_heading(doc, "8.2 结算后的玩家去向", level=2)
    add_bullets(
        doc,
        [
            "立即再来一局",
            "调整当前参数后重试",
            "查看个人记录榜",
            "查看总榜及与榜首差距",
        ],
    )

    add_heading(doc, "9. 数值设计", level=1)
    add_heading(doc, "9.1 底层属性", level=2)
    add_table(
        doc,
        ["属性", "含义", "对结果的主要影响"],
        [
            ("初速度", "起飞冲刺能力", "决定前段爆发速度和开局推进行为"),
            ("升力", "维持滞空与滑翔的能力", "决定中段飞行时间和滑翔距离"),
            ("稳定性", "抵御偏航、抖动和失控的能力", "决定成绩波动和发挥兑现率"),
            ("重量", "机体惯性与抗扰动能力", "决定抗风能力、下坠速度和开局速度"),
            ("阻力", "空气阻滞程度", "决定中后段减速速度"),
        ],
    )

    add_heading(doc, "9.2 参数映射关系", level=2)
    add_table(
        doc,
        ["来源", "主要提升", "主要削弱"],
        [
            ("机头尖锐度", "初速度、穿透感", "稳定性"),
            ("机翼宽度", "升力、滑翔能力", "初速度、抗扰动能力"),
            ("左右对称度", "稳定性、方向保持", "无直接削弱，但上限爆发较低"),
            ("尾翼上翘角度", "抬头能力、滑翔表现", "失速风险随角度增加"),
            ("纸张重量", "抗扰动、初段冲刺", "下坠速度和滑翔时长"),
        ],
    )

    add_heading(doc, "9.3 初版推荐数值原则", level=2)
    add_bullets(
        doc,
        [
            "玩家修改任一参数后，结算结果应有可感知变化，但不能产生失控级断层。",
            "均衡构筑应具备更高稳定出分能力，极端构筑应具备更高上限但更难驾驭。",
            "随机扰动只承担调味作用，建议总影响控制在5%至12%范围内。",
            "首版应避免唯一最优解，确保三种纸张都存在可被理解的使用场景。",
        ],
    )
    add_paragraph(doc, "推荐距离计算思路：最终距离 = 基础飞行能力 + 投掷修正 + 稳定发挥修正 + 随机扰动。该公式用于策划理解，不要求在代码中完全以显式公式呈现。", bold_prefix="推荐距离计算思路：")

    add_heading(doc, "10. 排行榜系统", level=1)
    add_heading(doc, "10.1 榜单类型", level=2)
    add_table(
        doc,
        ["榜单", "展示字段", "作用"],
        [
            ("总榜", "排名、昵称、最远距离、使用纸张、上传时间", "形成全服竞争目标"),
            ("个人记录榜", "历史最佳、最近10次成绩、最常用纸张、平均飞行距离", "帮助玩家复盘与优化"),
        ],
    )
    add_heading(doc, "10.2 业务规则", level=2)
    add_bullets(
        doc,
        [
            "每局结算后自动判断是否刷新个人最佳。",
            "若进入总榜，弹出恭喜上榜提示，强化成就反馈。",
            "排行榜支持按纸张类型筛选，鼓励玩家比较不同流派。",
            "个人记录榜默认展示最近10次成绩，便于玩家观察调参效果。",
        ],
    )
    add_heading(doc, "10.3 反作弊要求", level=2)
    add_bullets(
        doc,
        [
            "成绩上传时需提交纸张类型、参数配置、投掷输入结果和随机种子。",
            "关键得分逻辑建议由服务端复核，避免前端直接信任客户端距离结果。",
            "对异常高分建立阈值监控和人工复查机制。",
            "限制单设备或单账号短时间内异常频繁提交。",
        ],
    )

    add_heading(doc, "11. 页面结构与交互需求", level=1)
    add_table(
        doc,
        ["页面", "核心内容", "关键目标"],
        [
            ("首页", "开始游戏、玩法简介、个人最佳、排行榜入口", "快速进入游戏并建立目标感"),
            ("折纸页", "纸张选择、折纸步骤、参数滑块、性能预估", "完成构筑并理解参数作用"),
            ("投掷页", "蓄力、角度拖动、释放区提示", "形成一次有技术差异的投掷"),
            ("飞行页", "轨迹、距离数字、高度变化、风扰动表现", "强化紧张感和成绩期待"),
            ("结算页", "本次距离、纪录判断、失败原因、再来一局", "促成重开与优化"),
            ("排行榜页", "总榜、个人榜、昵称与成绩信息", "形成中长期留存动力"),
        ],
    )

    add_heading(doc, "12. 新手引导", level=1)
    add_paragraph(doc, "新手引导目标是在30秒内让玩家完成完整首局，同时理解四个最基本的因果关系。")
    add_bullets(
        doc,
        [
            "头更尖，飞得更快。",
            "翼更宽，飞得更久。",
            "更对称，更稳定。",
            "尾翼决定抬头还是下坠。",
        ],
    )
    add_paragraph(doc, "首局建议限制纸张为普通打印纸，并提供更明显的投掷最佳释放区提示，确保新玩家能快速看到一次完整的成功飞行。")

    add_heading(doc, "13. 视觉与音效方向", level=1)
    add_bullets(
        doc,
        [
            "视觉关键词：清爽、校园感、纸张手作感、轻竞技氛围。",
            "折纸阶段需要有明确折痕动画和纸张翻折反馈。",
            "飞行阶段距离数字应具备明显滚动感，制造冲纪录的紧张情绪。",
            "上榜和刷新纪录需有显著视觉强化与音效奖励。",
        ],
    )
    add_paragraph(doc, "音效建议包括纸张折叠声、投掷破风声、滑翔声、结算提示音和上榜强化音，整体风格保持轻快明亮。")

    add_heading(doc, "14. 数据埋点与成功指标", level=1)
    add_table(
        doc,
        ["指标", "目的"],
        [
            ("首局完成率", "验证新手引导是否顺畅"),
            ("前3局留存率", "验证首轮体验是否具备重复游玩动机"),
            ("平均单局时长", "验证节奏是否符合轻休闲定位"),
            ("纸张使用占比", "验证不同纸张是否都有存在感"),
            ("参数调整使用率", "验证玩家是否理解调参价值"),
            ("投掷力度与角度分布", "帮助优化投掷手感和最佳区间"),
            ("个人最佳刷新率", "衡量玩家是否持续获得进步反馈"),
            ("再来一局点击率", "衡量核心循环强度"),
            ("上榜率", "衡量排行榜门槛是否合理"),
        ],
    )

    add_heading(doc, "15. MVP范围与版本规划", level=1)
    add_heading(doc, "15.1 MVP首版必须包含", level=2)
    add_bullets(
        doc,
        [
            "3种纸张类型",
            "4个折纸核心参数",
            "1套完整投掷交互",
            "1套简化飞行模拟逻辑",
            "自由挑战模式",
            "总榜与个人记录榜",
            "新手引导",
            "基础埋点系统",
        ],
    )
    add_heading(doc, "15.2 首版明确不做", level=2)
    add_bullets(
        doc,
        [
            "复杂地图风场",
            "多场景障碍物",
            "飞行途中手动控制",
            "多玩法模式",
            "养成付费系统",
            "好友、公会等社交系统",
        ],
    )
    add_heading(doc, "15.3 后续可扩展方向", level=2)
    add_bullets(
        doc,
        [
            "限时挑战和周榜",
            "特殊纸张与主题赛季",
            "自定义外观皮肤",
            "更多飞行环境，如教室、操场、楼道等主题地图",
            "复盘系统和高手配置分享",
        ],
    )

    add_heading(doc, "16. 风险与注意事项", level=1)
    add_bullets(
        doc,
        [
            "如果随机扰动过强，玩家会将失败归因于运气，削弱调参与投掷的意义。",
            "如果参数反馈不清晰，玩家容易把折纸系统理解为无效装饰。",
            "如果榜单被作弊污染，会直接损伤核心竞争体验。",
            "如果首局流程过长，会降低网页端休闲玩家的进入意愿。",
        ],
    )

    add_heading(doc, "17. 结论", level=1)
    add_paragraph(doc, "《纸飞机冲榜》首版应围绕三件事构建体验闭环：一是参数化折纸带来的轻度构筑感，二是投掷动作带来的操作差异，三是排行榜带来的持续挑战欲。只要首版把这三项打磨清楚，就能形成短局高复玩的网页竞技产品雏形。")

    footer_section = doc.sections[-1]
    footer = footer_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("《纸飞机冲榜》正式PRD文档 V1.0")
    footer_run.font.name = "Microsoft YaHei"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    footer_run.font.size = Pt(9)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
